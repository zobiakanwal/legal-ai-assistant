# main.py
from fastapi.responses import StreamingResponse
from fastapi import FastAPI, HTTPException, Query
from fastapi.middleware.cors import CORSMiddleware
from typing import List
import os
import base64
from pathlib import Path
from docx import Document
from pydantic import BaseModel
import openai
import uuid
from datetime import datetime
from dotenv import load_dotenv
import json
import re
import logging
logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")

app = FastAPI()
GPT_MODEL = "gpt-4"

# CORS for frontend
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

BASE_DIR = Path(__file__).resolve().parent
TEMPLATES_DIR = BASE_DIR / "templates"

@app.get("/api/templates/{category}")
def list_templates(category: str):
    category_path = TEMPLATES_DIR / category
    if not category_path.exists() or not category_path.is_dir():
        raise HTTPException(status_code=404, detail="Category not found")
    files = [f.name for f in category_path.glob("*.docx")]
    return {"templates": files}

@app.get("/api/template")
def get_template(category: str = Query(...), name: str = Query(...)):
    file_path = TEMPLATES_DIR / category / name
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found")
    with open(file_path, "rb") as f:
        encoded = base64.b64encode(f.read()).decode("utf-8")
    return {"base64": encoded}

@app.get("/api/template/sections")
def get_template_sections(category: str = Query(...), name: str = Query(...)):
    file_path = TEMPLATES_DIR / category / name
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found")
    doc = Document(file_path)
    sections = [para.text.strip() for para in doc.paragraphs if para.text.strip().lower().startswith("template for")]
    if not sections:
        sections = ["Full Document"]
    return {"sections": sections}

load_dotenv()
openai.api_key = os.getenv("OPENAI_API_KEY")

class AIStartRequest(BaseModel):
    category: str
    subtype: str | None = None
    user_input: str

@app.post("/api/ai/start")
def start_ai_flow(data: AIStartRequest):
    try:
        if data.subtype:
            metadata_path = TEMPLATES_DIR / data.category / data.subtype / "metadata.json"
        else:
            metadata_path = TEMPLATES_DIR / data.category / "metadata.json"

        if not metadata_path.exists():
            raise HTTPException(status_code=500, detail="metadata.json not found")
        with open(metadata_path, "r", encoding="utf-8") as f:
            templates = json.load(f)
        if not templates:
            raise HTTPException(status_code=404, detail="No templates found.")
        for t in templates:
            if not all(k in t for k in ("title", "summary", "filename")):
                raise HTTPException(status_code=500, detail="Invalid metadata format")

        summaries_str = "\n\n".join([
            f"Title: {t['title']}\nSummary: {t['summary']}\nFilename: {t['filename']}"
            for t in templates
        ])

        selection_prompt = (
        "You are an expert legal assistant helping users find the most suitable document template.\n"
        f"A user described their issue as:\n\n{data.user_input.strip()}\n\n"
        "Here are the available legal document templates:\n\n"
        f"{summaries_str}\n\n"
        "Your task:\n"
        "— Select the single most suitable template.\n"
        "— Respond ONLY with the value of the 'filename' field (e.g. template_for_unreasonableness).\n"
        "— Do NOT explain your choice.\n"
        "— If multiple templates could apply, choose the best one based on the user's description.\n"
        "— If unsure, pick the closest reasonable match anyway — do NOT say 'none match'."
    )

        try:
            response = openai.chat.completions.create(
                model= GPT_MODEL,
                messages=[
                    {"role": "system", "content": "You are a helpful legal assistant..."},
                    {"role": "user", "content": selection_prompt},
                ],
                temperature=0.3,
                max_tokens=60,
            )
        
        except openai.OpenAIError as e:
            logging.error(f"OpenAI API error during template selection: {str(e)}")
            raise HTTPException(status_code=500, detail="AI service error")

        selected_filename = response.choices[0].message.content.strip()
        logging.info(f"Selected Filename from GPT: {selected_filename}")

        matched_template = next(
            (t for t in templates if t["filename"].strip().lower() == selected_filename.lower()),
            None
        )

        if not matched_template:
            raise HTTPException(status_code=404, detail="Template match not found in metadata")

        doc_filename = f"{selected_filename}.docx"
        logging.info(f"Selected Template File: {doc_filename}")

        if data.subtype:
            doc_path = TEMPLATES_DIR / data.category / data.subtype / doc_filename
        else:
            doc_path = TEMPLATES_DIR / data.category / doc_filename

        if not doc_path.exists():
            raise HTTPException(status_code=404, detail="Selected template file not found")

        doc = Document(doc_path)
        full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        question_prompt = (
            "You are a professional legal assistant helping a user complete a legal document.\n"
            "You are given the full text of the template. Read it carefully and identify all placeholders or gaps that must be completed by the user (e.g. [insert full address], empty lines, bullet point options, or areas left blank for details).\n\n"
            "Ask questions one at a time to gather the exact information needed to fill in these blanks. Start with the most essential or obvious missing fields.\n"
            "Make each question clear, simple, and specific — just like you're guiding someone through a form.\n"
            "If there are multiple options in a section (e.g. a, b, c), ask follow-up questions to help the user choose the correct one.\n"
            "Do not explain the document. Just act like a legal assistant who knows what details are needed and asks for them naturally, one by one.\n"
            "Avoid asking for contact info unless the template explicitly requires it.\n\n"
            "Once the necessary information has been collected, the document will be auto-completed and downloaded by the user."
        )

        try:
            q_response = openai.chat.completions.create(
                model= GPT_MODEL,
                messages=[
                    {"role": "system", "content": question_prompt},
                    {"role": "user", "content": full_text},
                ],
                temperature=0.3,
                max_tokens=150,
            )
        except openai.OpenAIError as e:
            logging.error(f"OpenAI API error during first question generation: {str(e)}")
            raise HTTPException(status_code=500, detail="AI service error")

        question = q_response.choices[0].message.content.strip()
        return {
        "question": question,
        "filename": f"{data.subtype}/{selected_filename}" if data.subtype else selected_filename
    }

    except Exception as e:
        logging.exception("Unexpected server error")
        raise HTTPException(status_code=500, detail=str(e))
    
class AINextRequest(BaseModel):
    category: str
    filename: str
    messages: List[dict]

@app.post("/api/ai/next")
def ai_next_question(data: AINextRequest):
    try:
        subtype_path = data.filename.split("/", 1)
        if len(subtype_path) == 2:
            subfolder, filename = subtype_path
            file_path = TEMPLATES_DIR / data.category / subfolder / f"{filename}.docx"
        else:
            file_path = TEMPLATES_DIR / data.category / f"{data.filename}.docx"

        if not file_path.exists():
            raise HTTPException(status_code=404, detail="Template not found")

        doc = Document(file_path)
        template_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

        system_prompt = (
            "You are a professional legal assistant continuing a session to help a user complete a legal document.\n"
            "You have access to the full document template and the conversation history.\n"
            "Identify any remaining placeholders (like [insert...], blank lines, bullet point choices, or missing details).\n"
            "Ask ONE specific, clear question at a time to gather that missing information.\n"
            "If all placeholders are filled, respond with __COMPLETE__ to signal the document is ready for generation.\n"
            "Do not explain or summarize the document — focus only on gathering the required inputs naturally and efficiently."
        )

        response = openai.chat.completions.create(
            model=GPT_MODEL,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": f"This is the template:\n\n{template_text}"},
                *data.messages
            ],
            temperature=0.3,
            max_tokens=300
        )

        reply = response.choices[0].message.content.strip()
        return {"reply": reply}

    except openai.OpenAIError as e:
        logging.error(f"OpenAI API error during Q&A: {str(e)}")
        raise HTTPException(status_code=500, detail="AI service error")

    except Exception as e:
        logging.exception("Unexpected server error during /api/ai/next")
        raise HTTPException(status_code=500, detail=str(e))

class AICompleteRequest(BaseModel):
    category: str
    filename: str
    messages: List[dict]

def extract_answers(messages: List[dict]):
    return [messages[i]["content"] for i in range(1, len(messages), 2) if messages[i]["role"] == "user"]

@app.get("/api/categories")
def list_categories():
    categories = {}
    for cat in TEMPLATES_DIR.iterdir():
        if not cat.is_dir(): continue
        subtypes = [sub.name for sub in cat.iterdir() if sub.is_dir() and ((sub / "metadata.json").exists() or any(f.suffix == ".docx" for f in sub.glob("*.docx")))]
        if subtypes or any(f.suffix == ".docx" for f in cat.glob("*.docx")) or (cat / "metadata.json").exists():
            categories[cat.name] = subtypes
    return categories

@app.post("/api/ai/complete")
def complete_template(data: AICompleteRequest):
    filename = data.filename.replace(".docx", "")
    file_path = TEMPLATES_DIR / data.category / f"{filename}.docx"
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="Template not found")
    doc = Document(file_path)

    # GPT-driven fill‑in logic
    raw_template = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    chat_log = "\n".join(f"{m['role']}: {m['content']}" for m in data.messages)

    fill_prompt = f"""
You are a professional legal assistant.
Fill every gap in the document below (brackets, underlines, dotted blanks) using only the conversation data.

DOCUMENT TEMPLATE:
{raw_template}

CONVERSATION:
{chat_log}

Return JSON with:
  nextQuestion: <string or null>
  filledDocument: <complete text or null>
"""

    resp = openai.chat.completions.create(
        model=GPT_MODEL,
        messages=[
            {"role":"system","content":"You fill in and ask for missing info."},
            {"role":"user","content": fill_prompt}
        ],
        temperature=0.2,
        max_tokens=2000
    )
    result = json.loads(resp.choices[0].message.content)

    if result["nextQuestion"]:
        return {"nextQuestion": result["nextQuestion"]}

    # build final .docx from result["filledDocument"]
    from io import BytesIO
    out_doc = Document()
    for line in result["filledDocument"].split("\n"):
        out_doc.add_paragraph(line)
    buf = BytesIO()
    out_doc.save(buf)
    buf.seek(0)
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition":f"attachment; filename=filled_{filename}.docx"}
    )
